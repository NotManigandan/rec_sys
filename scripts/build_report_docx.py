from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
REPORT_JSON = ROOT / "artifacts" / "model_report.json"
OUTPUT_PATH = ROOT / "docs" / "recommendation_methods_report.docx"


def add_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)


def format_metric(value: float) -> str:
    return f"{value:.4f}"


def main() -> None:
    report = json.loads(REPORT_JSON.read_text())

    document = Document()
    document.add_heading("Recommendation System Methods and Evaluation Report", level=0)

    document.add_heading("1. Overview", level=1)
    add_paragraphs(
        document,
        [
            "This report summarizes the recommendation methods currently implemented for the synthetic recommendation-system dataset in rec_system/ and explains how model performance is measured.",
            "The dataset contains 2,000 users, 6 items, 3 reported ratings per user in reported.csv, and full hidden ground-truth ratings in actual.csv.",
            "The training setup uses the partially observed ratings in reported.csv to learn a recommender. The learned model is then evaluated on the 3 unseen items for each user using the hidden truth in actual.csv.",
            "This setup is useful because it provides known ground truth, making it possible to measure whether the recommender is learning users' actual latent preferences rather than only matching observed ratings.",
        ],
    )

    document.add_heading("2. Methods", level=1)

    document.add_heading("2.1 Popularity Baseline", level=2)
    add_paragraphs(
        document,
        [
            "The popularity baseline is the simplest method. It does not personalize recommendations by user. Instead, it recommends items with the highest average reported rating across the full dataset.",
            "Strengths: very simple, fast to compute, and useful as a sanity-check baseline.",
            "Weaknesses: ignores user-specific preferences and tends to over-recommend globally popular items even when they are not optimal for a specific user.",
        ],
    )

    document.add_heading("2.2 BPR Matrix Factorization (bpr_mf)", level=2)
    add_paragraphs(
        document,
        [
            "This is the main baseline and the strongest overall model in the current benchmark.",
            "BPR stands for Bayesian Personalized Ranking. Each user gets a learned embedding vector, each item gets a learned embedding vector, and the user-item score is computed mainly through a dot product plus user and item bias terms.",
            "The model is trained using pairwise ranking comparisons. If a user gave one observed item a higher rating than another observed item, the model is trained to score the higher-rated item above the lower-rated item.",
            "Strengths: directly optimized for ranking, simple, interpretable, and well matched to collaborative-filtering data without rich side features.",
            "Weaknesses: less expressive than deeper nonlinear models and relies mostly on interaction structure rather than metadata.",
        ],
    )

    document.add_heading("2.3 Neural Collaborative Filtering (neural_cf)", level=2)
    add_paragraphs(
        document,
        [
            "This method is a more flexible nonlinear version of collaborative filtering.",
            "Like matrix factorization, it learns user and item embeddings. However, instead of using only a dot product, it concatenates the embeddings and passes them through a multilayer perceptron to produce a score.",
            "Strengths: more expressive than standard matrix factorization, can model more complex user-item interactions, and serves as a stronger nonlinear comparison model.",
            "Weaknesses: can overfit more easily on small datasets, is less interpretable, and the extra complexity does not always improve top recommendation quality.",
        ],
    )

    document.add_heading("2.4 Two-Tower Model (two_tower)", level=2)
    add_paragraphs(
        document,
        [
            "The two-tower model separates the user representation and item representation into two different encoders.",
            "In the current implementation, the user tower converts the user's reported rating history into a user vector, while the item tower uses a learned item embedding. The final score is the similarity between the user vector and the item vector.",
            "Strengths: good architecture for large-scale retrieval, easy to extend with side features, and useful if the project later includes richer metadata or cold-start settings.",
            "Weaknesses: less effective on this dataset because there are only 6 items, only 3 reported ratings per user, and no rich side features.",
        ],
    )

    document.add_heading("3. Evaluation Metrics", level=1)
    add_paragraphs(
        document,
        [
            "The evaluation process ranks the 3 unseen items for each user and compares that ranking against the user's hidden true ratings in actual.csv.",
        ],
    )

    document.add_heading("3.1 Top-1 Accuracy", level=2)
    add_paragraphs(
        document,
        [
            "Top-1 accuracy measures whether the model's highest-ranked unseen item is actually one of the user's best unseen items.",
            "Interpretation: 1.0 means the model always places a truly best hidden item at the top, while 0.0 means it never does.",
        ],
    )

    document.add_heading("3.2 NDCG@3", level=2)
    add_paragraphs(
        document,
        [
            "NDCG@3 evaluates the quality of the full ranking of the 3 hidden items. Higher-relevance items are rewarded more, and items ranked earlier receive more credit.",
            "Interpretation: 1.0 means a perfect ranking, and values closer to 1.0 indicate that the predicted ordering is close to the ideal hidden ordering.",
        ],
    )

    document.add_heading("3.3 Pairwise Accuracy", level=2)
    add_paragraphs(
        document,
        [
            "Pairwise accuracy checks whether the model correctly orders pairs of hidden items whenever the two items have different true ratings.",
            "Interpretation: 1.0 means the model always orders unequal pairs correctly, while values near 0.5 are much closer to random guessing.",
        ],
    )

    document.add_heading("3.4 MRR", level=2)
    add_paragraphs(
        document,
        [
            "MRR, or Mean Reciprocal Rank, measures how early the best hidden item appears in the ranked list.",
            "If the best item is ranked first, the score is 1.0. If it is ranked second, the score is 0.5. If it is ranked third, the score is 0.333. Higher values mean the best item usually appears near the top.",
        ],
    )

    document.add_heading("4. Results", level=1)
    document.add_paragraph("Table 1. Overall model performance on hidden-item ranking.")

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Method"
    header_cells[1].text = "Top-1 Accuracy"
    header_cells[2].text = "NDCG@3"
    header_cells[3].text = "Pairwise Accuracy"
    header_cells[4].text = "MRR"

    row_order = [
        ("popularity", "Popularity Baseline"),
        ("bpr_mf", "BPR Matrix Factorization (bpr_mf)"),
        ("neural_cf", "Neural Collaborative Filtering (neural_cf)"),
        ("two_tower", "Two-Tower Model (two_tower)"),
    ]
    for key, label in row_order:
        metrics = report[key]["metrics"]
        row = table.add_row().cells
        row[0].text = label
        row[1].text = format_metric(metrics["top1_accuracy"])
        row[2].text = format_metric(metrics["ndcg@3"])
        row[3].text = format_metric(metrics["pairwise_accuracy"])
        row[4].text = format_metric(metrics["mrr"])

    document.add_heading("5. Interpretation of Results", level=1)
    add_paragraphs(
        document,
        [
            "The main conclusions are that bpr_mf is the strongest overall model on this dataset, neural_cf is competitive and slightly better on pairwise ordering, two_tower underperforms in this small feature-poor setting, and the learned personalized models clearly outperform the non-personalized popularity baseline on most metrics.",
            "BPR performs best here because the dataset is small, the task is inherently ranking-based, and there are no rich user or item features that would justify a more complex architecture.",
            "The two-tower model underperforms because it is designed for richer feature settings, while the current dataset has only 3 observed ratings per user and only 6 items total.",
        ],
    )

    document.add_heading("6. Recommendation", level=1)
    add_paragraphs(
        document,
        [
            "For the next phase of the project, the recommended modeling plan is to use bpr_mf as the primary baseline recommender, keep neural_cf as a stronger nonlinear comparison model, use the popularity model as the non-personalized sanity-check baseline, and avoid prioritizing the two-tower model unless the dataset is expanded with richer user or item features.",
            "This creates a clean story for later robustness experiments: a clean baseline recommender, an attacked version, and a defended or robustified version.",
        ],
    )

    document.add_heading("7. Limitations", level=1)
    add_paragraphs(
        document,
        [
            "The current results should be interpreted in the context of the dataset design: only 6 items are available, each user reports only 3 observed ratings, the dataset is synthetic, and no timestamps, text, or metadata features are included.",
            "Because of this, the current benchmark is best understood as a controlled proof of concept rather than a realistic production recommendation problem.",
        ],
    )

    document.add_heading("8. Appendix: One-Sentence Summary of Each Method", level=1)
    add_paragraphs(
        document,
        [
            "Popularity Baseline: recommend globally popular items to everyone.",
            "BPR Matrix Factorization: learn user and item latent vectors and rank higher-rated observed items above lower-rated observed items.",
            "Neural Collaborative Filtering: learn user and item embeddings and use an MLP to model nonlinear interactions.",
            "Two-Tower Model: encode user history and item identity separately and score them by similarity.",
        ],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
