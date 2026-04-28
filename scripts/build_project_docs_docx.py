from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"


def render_markdown_to_docx(source_path: Path, output_path: Path) -> None:
    document = Document()

    lines = source_path.read_text(encoding="utf-8").splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
            continue

        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        numbered_prefix, dot, remainder = stripped.partition(". ")
        if dot and numbered_prefix.isdigit():
            document.add_paragraph(remainder.strip(), style="List Number")
            continue

        document.add_paragraph(stripped)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(f"wrote {output_path}")


def main() -> None:
    jobs = [
        (DOCS_DIR / "project_description.md", DOCS_DIR / "project_description.docx"),
        (DOCS_DIR / "poster_outline.md", DOCS_DIR / "poster_outline.docx"),
    ]
    for source_path, output_path in jobs:
        render_markdown_to_docx(source_path, output_path)


if __name__ == "__main__":
    main()
